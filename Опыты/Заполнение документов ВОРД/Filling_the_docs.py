from docx import Document
from docx import styles
from docx import shared
from docx.enum.style import WD_STYLE_TYPE
import re
from BD_sample import form_data, dadata_data #импортируем из соседнего файла наши словари

class ResponsiblePerson:
    def __init__(self, doc_name):
        self.document = Document(doc_name)
        self.add_new_style()

    def add_new_style(self):
        obj_styles = self.document.styles
        obj_charstyle = obj_styles.add_style('MainStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = shared.Pt(14)
        obj_font.name = 'Cambria'

    def fill_organisation_name(self, organisation):
        self.document.paragraphs[0].clear()
        self.document.paragraphs[1].clear()
        self.document.paragraphs[0].add_run(organisation, style='MainStyle').bold = True

    def fill_company_adress(self, company_adress):
        self.document.paragraphs[2].clear()
        self.document.paragraphs[2].add_run(company_adress).bold = True

    def fill_ogrn_inn_kpp(self, ogrn, inn, kpp):
        # записываем текст параграфа в переменную
        par_text = self.document.paragraphs[3].text

        # убираем текст параграфа из документа, чтобы переписать его с изменениями
        self.document.paragraphs[3].clear()

        # меняем старый текст (par_text), который собираемся записать в документ
        par_text = re.sub('ОГРН_+', f'ОГРН {ogrn}', par_text)
        par_text = re.sub('ИНН _+', f'ИНН {inn}', par_text)
        par_text = re.sub('КПП _+', f'КПП {kpp}', par_text)

        # записываем изменённый текст обратно в документ
        self.document.paragraphs[3].add_run(par_text).bold = True




    def fill_city_and_date(self, rowcity, date):
        # записываем текст параграфа в переменную
        par_text = self.document.paragraphs[8].text

        # убираем текст параграфа из документа, чтобы переписать его с изменениями
        self.document.paragraphs[8].clear()

        clean_city = rowcity.split(', ')[1]
        day, month, year = date.split('-')

        # меняем старый текст (par_text), который собираемся записать в документ
        par_text = re.sub('^_+', clean_city, par_text)
        par_text = re.sub('          _____\\.', f' {day}.', par_text)
        par_text = re.sub('\\._____\\.', f' {month} ', par_text)
        par_text = re.sub(' _+$', f' {year} ', par_text)
        # записываем изменённый текст обратно в документ
        self.document.paragraphs[8].add_run(par_text).bold = True

    def fill_company_full_name(self, company_full_name):
        # записываем текст параграфа в переменную
        par_text = self.document.paragraphs[10].text
        print(par_text) #через этот принт мы смотрим какой параграф с каким номером в исходном документе

        # убираем текст параграфа из документа, чтобы переписать его с изменениями
        self.document.paragraphs[10].clear()

        # меняем старый текст (par_text), который собираемся записать в документ
        par_text = re.sub(' _+',' '+company_full_name, par_text)

           # записываем изменённый текст обратно в документ
        self.document.paragraphs[10].add_run(par_text)

    def fill_CEO_signature(self, CEO_position, CEO_FIO):
        # записываем текст параграфа в переменную
        par_text = self.document.paragraphs[18].text

        # убираем текст параграфа из документа, чтобы переписать его с изменениями
        self.document.paragraphs[18].clear()

        # меняем старый текст (par_text), который собираемся записать в документ
        par_text = re.sub('^_+', f'{CEO_position.capitalize()}', par_text)
        par_text = re.sub('_+$', CEO_FIO, par_text)
        #par_text = re.sub('КПП _+', f'КПП {kpp}', par_text)

        # записываем изменённый текст обратно в документ
        run = self.document.paragraphs[18].add_run(par_text)
        run.bold = True

    def __str__(self):
        result = ''
        for par in self.document.paragraphs:
            result += '\n' + par.text
        return result

doc_name = 'Document_templates/1. Приказ о назначении ответственного за обраотку ПДн, должностные обязанности.docx'


responsible_person = ResponsiblePerson(doc_name)
responsible_person.fill_organisation_name(dadata_data ['short_name_company'])
responsible_person.fill_ogrn_inn_kpp(dadata_data ['OGRN_company'], dadata_data['INN_company'],dadata_data['kpp_company'])
responsible_person.fill_city_and_date(dadata_data['address_company'], form_data['sign_date'])
responsible_person.fill_company_adress(dadata_data['address_company'])
responsible_person.fill_company_full_name(dadata_data ['full_name_company'])
responsible_person.fill_CEO_signature(dadata_data ['CEO'],dadata_data['CEO_name'])


# print(responsible_person)
responsible_person.document.save('Output_docs/demo.docx')

#print(responsible_person.document.paragraphs[18].text) #через этот принт мы смотрим
# какие строки(параграфы) по счету в исходном документы для корректировки и замены